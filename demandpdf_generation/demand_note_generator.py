"""
Demand Note Generator (JSON-based)
---------------------------------
Professional legal demand letter generator (DOCX)

Implements requirements 1–10 exactly
Uses python-docx with precise spacing control
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =====================================================
# JSON LOADER
# =====================================================

def load_metadata_from_json(json_path: str) -> Dict[str, Any]:
    path = Path(json_path)
    if not path.exists():
        raise FileNotFoundError(f"Metadata JSON not found: {path}")

    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    for key in ["demand_creation_date", "date_of_loss"]:
        if key in data and isinstance(data[key], str):
            try:
                data[key] = datetime.fromisoformat(data[key])
            except ValueError:
                pass

    return data


# =====================================================
# DEMAND NOTE GENERATOR
# =====================================================

class DemandNoteGenerator:
    FONT_NAME = "Franklin Gothic Book"
    FONT_SIZE = 12

    def __init__(self):
        self.doc = None

    # -------------------------------------------------
    # Utilities
    # -------------------------------------------------

    def _set_font(self, run, size=None, bold=False, underline=False):
        run.font.name = self.FONT_NAME
        run.font.size = Pt(size or self.FONT_SIZE)
        run.font.bold = bold
        run.font.underline = underline

    def _keep_with_next(self, paragraph):
        pPr = paragraph._element.get_or_add_pPr()
        pPr.append(OxmlElement("w:keepLines"))
        pPr.append(OxmlElement("w:keepNext"))

    def _tight_paragraph(self, paragraph, before=0, after=0):
        pPr = paragraph._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), str(before))
        spacing.set(qn("w:after"), str(after))
        pPr.append(spacing)

    def _add_border(self, paragraph):
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for side in ["top", "left", "bottom", "right"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "24")
            el.set(qn("w:color"), "000000")
            pBdr.append(el)
        pPr.append(pBdr)

    # -------------------------------------------------
    # Document Builder
    # -------------------------------------------------

    def create_demand_note(self, m: Dict[str, Any], output_path: str) -> str:
        self.doc = Document()

        style = self.doc.styles["Normal"]
        style.font.name = self.FONT_NAME
        style.font.size = Pt(self.FONT_SIZE)

        self._add_logo_and_date(m)
        self._add_insurance_block(m)
        self._add_title(m)
        self._add_client_info(m)
        self._add_settlement_notice()
        self._add_salutation(m)
        self._add_incident(m)
        self._add_medical_summary(m)
        self._add_compensation(m)

        self.doc.save(output_path)
        return output_path

 

    def _add_logo_and_date(self, m):
        logo = m.get("logo_path")
        if logo and Path(logo).exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(logo, width=Inches(1))
            self._tight_paragraph(p, after=40)
            self._keep_with_next(p)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_text = (
            m["demand_creation_date"].strftime("%B %d, %Y")
            if isinstance(m.get("demand_creation_date"), datetime)
            else m.get("demand_creation_date", "")
        )
        r = p.add_run(date_text)
        self._set_font(r, size=14, bold=True, underline=True)
        self._tight_paragraph(p, after=80)
        self._keep_with_next(p)

    def _add_insurance_block(self, m):
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = False

        left = table.rows[0].cells[0].paragraphs[0]
        for line in [
            m.get("insurance_name"),
            f"Attn: {m.get('claim_number')}" if m.get("claim_number") else None,
            m.get("insurance_address"),
            f"Tel: {m.get('insurance_telephone')}" if m.get("insurance_telephone") else None,
            f"Fax: {m.get('insurance_fax')}" if m.get("insurance_fax") else None,
        ]:
            if line:
                r = left.add_run(line + "\n")
                self._set_font(r)

        right = table.rows[0].cells[1].paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = right.add_run("Sent Via Certified U.S. Mail\nFacsimile")
        self._set_font(r)

        self._keep_with_next(left)

    def _add_title(self, m):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"*** {m.get('accident_type','ACCIDENT').upper()} POLICY LIMIT DEMAND ***")
        self._set_font(r, size=18, bold=True, underline=True)
        self._tight_paragraph(p, before=80, after=80)
        self._keep_with_next(p)

    def _add_client_info(self, m):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fields = [
            ("Client", m.get("client_name")),
            ("Date of Loss",
             m["date_of_loss"].strftime("%m/%d/%Y")
             if isinstance(m.get("date_of_loss"), datetime)
             else m.get("date_of_loss")),
            ("Your Insured", m.get("defendant_name")),
            ("Claim #", m.get("claim_number")),
            ("Adjuster", m.get("defendant_adjuster")),
        ]

        for label, value in fields:
            if value:
                r = p.add_run(f"{label}: {value}\n")
                self._set_font(r)

        self._tight_paragraph(p, after=80)
        self._keep_with_next(p)

    def _add_settlement_notice(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            "THIS LETTER IS FOR THE PURPOSE OF SETTLEMENT ONLY;\n"
            "IT IS NOT TO BE USED AS EVIDENCE OF ANY KIND.\n"
            "(CALIFORNIA EVIDENCE CODE SECTIONS 1152-54, ET SEQ.)\n"
            "THE CONTENTS OF THIS DEMAND SHALL NOT BE REPRODUCED\n"
            "OR REDISTRIBUTED FOR ANY PURPOSE."
        )
        self._set_font(r, bold=True, underline=True)
        self._tight_paragraph(p, before=80, after=80)
        self._keep_with_next(p)

    def _add_salutation(self, m):
        p = self.doc.add_paragraph()
        r = p.add_run("To Whom It May Concern:")
        self._set_font(r)
        self._tight_paragraph(p, after=20)

        intro = self.doc.add_paragraph()

        r1 = intro.add_run("This office represents ")
        self._set_font(r1)

        display = f"{m.get('title','')} {m.get('client_name')}"
        short = f"{m.get('title','')} {m.get('client_last_name')}"

        r2 = intro.add_run(f"{display} (“{short}”)")
        self._set_font(r2, bold=True)

        r3 = intro.add_run(
            " in the above-referenced incident involving your insured. "
            "We hereby extend this formal offer of settlement as set forth herein."
        )
        self._set_font(r3)

    def _add_incident(self, m):
        head = self.doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = head.add_run("THE INCIDENT")
        self._set_font(r, size=16, bold=True, underline=True)
        self._tight_paragraph(head, before=80, after=40)
        self._keep_with_next(head)

        if m.get("incident_summary"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(m["incident_summary"])
            self._set_font(r)

    # =====================================================
    # MEDICAL SUMMARY
    # =====================================================

    def _add_medical_summary(self, m):
        records = m.get("medical_records")
        if not records:
            return

    # 🔴 FORCE PAGE 2 START
        self.doc.add_page_break()

        head = self.doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = head.add_run("SUMMARY OF MEDICAL CARE & INJURIES")
        self._set_font(r, size=16, bold=True, underline=True)
        self._tight_paragraph(head, before=80, after=80)

        for rec in records:
            if rec.get("summary"):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                r = p.add_run(rec["summary"])
                self._set_font(r)
                self._tight_paragraph(p, after=60)

            if rec.get("image_path") and Path(rec["image_path"]).exists():
                label = self.doc.add_paragraph()
                r = label.add_run(rec.get("image_reference", "IMAGE"))
                self._set_font(r, bold=True)
                self._tight_paragraph(label, after=20)

                img_p = self.doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.add_run().add_picture(rec["image_path"], width=Inches(4.5))
                self._add_border(img_p)
                self._tight_paragraph(img_p, after=80)


    # =====================================================
    # CLAIM FOR COMPENSATION (NEW PAGE)
    # =====================================================

    def _add_compensation(self, m):
        self.doc.add_page_break()

        title = m.get("title", "")
        last = m.get("client_last_name", "")

        head = self.doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = head.add_run("CLAIM FOR COMPENSATION")
        self._set_font(r, size=16, bold=True, underline=True)
        self._tight_paragraph(head, after=120)

        self._left_para(
            f"{title}. {last} is entitled to full and fair compensation for losses incurred "
            "in this accident. Relevant California law provides as follows:"
        )

        self._center_underline("Civil Code § 3333.  Torts in General")

        self._center_para(
            "Breach of obligation other than contract. For the breach of an obligation not "
            "arising from contract, the measure of damages, except where otherwise expressly "
            "provided by this code, is the amount which will compensate for all the detriment "
            "proximately caused thereby, whether it could have been anticipated or not."
        )

        self._left_para(
            f"{title}. {last} is also entitled to recover for all the pain and suffering "
            "experienced as a result of this accident:"
        )

        self._center_para(
            "Pain and suffering is a unitary concept, encompassing all of the physical and "
            "emotional trauma occasioned by an injury. Plaintiff is entitled to compensatory "
            "damages for all physical pain suffered . . . and also for resulting fright, "
            "nervousness, anxiety, worry, mortification, shock, humiliation, indignity, "
            "embarrassment, apprehension, terror, or ordeal. "
            "Capelouto v. Kaiser Foundation Hospitals (1972) 7 Cal.App.3d 889, 893-894."
        )

        self._left_para("The applicable jury instruction provides, in pertinent part, as follows:")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("CACI 3905A")
        self._set_font(r, size=16, bold=True)
        self._tight_paragraph(p, after=80)

        self._center_para(
            f"{title}. {last} is also entitled to recover for all past and future noneconomic "
            "damages caused by your insured’s negligence, including physical pain, mental "
            "suffering, loss of enjoyment of life, disfigurement, physical impairment, "
            "inconvenience, grief, anxiety, humiliation, and emotional distress. "
            "(CACI 3905A, Physical Pain, Mental Suffering, and Emotional Distress)"
        )

    # -------------------------------------------------
    # Helpers for compensation
    # -------------------------------------------------

    def _left_para(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        self._set_font(r)
        self._tight_paragraph(p, after=80)

    def _center_para(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        self._set_font(r)
        self._tight_paragraph(p, after=80)

    def _center_underline(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        self._set_font(r, underline=True)
        self._tight_paragraph(p, after=80)


# =====================================================
# MAIN
# =====================================================

if __name__ == "__main__":
    metadata = load_metadata_from_json("demand_metadata.json")

    generator = DemandNoteGenerator()
    output = generator.create_demand_note(
        metadata,
        "demand_note_sample.docx"
    )

    print(f"✅ Demand note created: {output}")

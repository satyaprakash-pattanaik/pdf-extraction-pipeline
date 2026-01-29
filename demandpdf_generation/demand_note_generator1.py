"""
Demand Note Generator (Pure python-docx)
---------------------------------------
Professional legal demand letter generator (DOCX)
Matches law-firm litigation templates
"""

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =====================================================
# JSON LOADER
# =====================================================

def load_metadata(json_path: str) -> dict:
    path = Path(json_path)
    if not path.exists():
        raise FileNotFoundError(f"Metadata JSON not found: {path}")

    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    for key in ["demand_creation_date", "date_of_accident"]:
        if isinstance(data.get(key), str):
            try:
                data[key] = datetime.fromisoformat(data[key])
            except ValueError:
                pass

    data["claim_number"] = data.get("client_claim_number")
    data["date_of_loss"] = data.get("date_of_accident")
    data["defendant_adjuster"] = data.get("defendant_adjuster_name")

    if isinstance(data.get("client_name"), list):
        data["client_display_name"] = ", ".join(data["client_name"])
        data["client_last_name"] = data["client_name"][0].split()[-1]
    else:
        data["client_display_name"] = data.get("client_name", "")
        data["client_last_name"] = data["client_display_name"].split()[-1]

    return data


# =====================================================
# DEMAND NOTE GENERATOR
# =====================================================

class DemandNoteGenerator:
    FONT = "Franklin Gothic Book"
    SIZE = 12

    def __init__(self):
        self.doc = Document()
        style = self.doc.styles["Normal"]
        style.font.name = self.FONT
        style.font.size = Pt(self.SIZE)

    # -------------------------------------------------
    # Helpers
    # -------------------------------------------------

    def _set_font(self, run, size=None, bold=False, underline=False):
        run.font.name = self.FONT
        run.font.size = Pt(size or self.SIZE)
        run.font.bold = bold
        run.font.underline = underline

    def _spacing(self, p, before=0, after=0):
        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), str(before))
        spacing.set(qn("w:after"), str(after))
        pPr.append(spacing)

    def _add_border(self, p):
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for side in ["top", "left", "bottom", "right"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "24")
            el.set(qn("w:color"), "000000")
            pBdr.append(el)
        pPr.append(pBdr)

    # -------------------------------------------------
    # Sections
    # -------------------------------------------------

    def add_logo_and_date(self, m):
        if m.get("logo_path") and Path(m["logo_path"]).exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(m["logo_path"], width=Inches(2))
            self._spacing(p, after=60)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(m["demand_creation_date"].strftime("%B %d, %Y"))
        self._set_font(r, size=14, bold=True, underline=True)
        self._spacing(p, after=80)

    def add_insurance_block(self, m):
        table = self.doc.add_table(rows=1, cols=2)
        left = table.rows[0].cells[0].paragraphs[0]

        for line in [
            m.get("insurance_name"),
            f"Attn: {m.get('claim_number')}",
            m.get("insurance_company_address"),
            f"Tel: {m.get('insurance_telephone')}",
            f"Fax: {m.get('insurance_fax')}",
        ]:
            if line:
                r = left.add_run(line + "\n")
                self._set_font(r)

        right = table.rows[0].cells[1].paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = right.add_run("Sent Via Certified U.S. Mail\nFacsimile")
        self._set_font(r)

    def add_title(self, m):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"*** {m.get('accident_type','ACCIDENT').upper()} POLICY LIMIT DEMAND ***")
        self._set_font(r, size=18, bold=True, underline=True)
        self._spacing(p, before=80, after=80)

    def add_client_info(self, m):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for label, value in [
            ("Client", m["client_display_name"]),
            ("Date of Loss", m["date_of_loss"].strftime("%m/%d/%Y")),
            ("Your Insured", m.get("defendant_name")),
            ("Claim #", m.get("claim_number")),
            ("Adjuster", m.get("defendant_adjuster")),
        ]:
            if value:
                r = p.add_run(f"{label}: {value}\n")
                self._set_font(r)

        self._spacing(p, after=80)

    def add_settlement_notice(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            "THIS LETTER IS FOR THE PURPOSE OF SETTLEMENT ONLY;\n"
            "IT IS NOT TO BE USED AS EVIDENCE OF ANY KIND.\n"
            "(CALIFORNIA EVIDENCE CODE SECTIONS 1152-54, ET SEQ.)\n"
            "THE CONTENTS OF THIS DEMAND SHALL NOT BE REPRODUCED\n"
            "OR REDISTRIBUTED FOR ANY PURPOSE."
        )
        self._set_font(r, bold=True, underline=True)
        self._spacing(p, before=80, after=80)

    def add_salutation_and_intro(self, m):
        p = self.doc.add_paragraph()
        self._set_font(p.add_run("To Whom It May Concern:"))
        self._spacing(p, after=20)

        intro = self.doc.add_paragraph()
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self._set_font(intro.add_run("This office represents "))
        self._set_font(
            intro.add_run(
                f"{m.get('title','')} {m['client_display_name']} "
                f"(“{m.get('title','')} {m['client_last_name']}”)"
            ),
            bold=True
        )
        self._set_font(
            intro.add_run(
                " in the above-referenced incident involving your insured. "
                "We hereby extend this formal offer of settlement as set forth herein."
            )
        )
        self._spacing(intro, after=80)

    def add_incident_and_medical(self, m):
        head = self.doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_font(head.add_run("THE INCIDENT"), size=16, bold=True, underline=True)
        self._spacing(head, before=80, after=40)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._set_font(p.add_run(m.get("incident_summary", "")))
        self._spacing(p, after=60)

        if not m.get("medical_records"):
            return

        head = self.doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_font(head.add_run("SUMMARY OF MEDICAL CARE & INJURIES"), size=16, bold=True, underline=True)
        self._spacing(head, after=60)

        for rec in m["medical_records"]:
            if rec.get("summary"):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._set_font(p.add_run(rec["summary"]))
                self._spacing(p, after=60)

            if rec.get("image_path") and Path(rec["image_path"]).exists():
                img = self.doc.add_paragraph()
                img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img.add_run().add_picture(rec["image_path"], width=Inches(4.5))
                self._add_border(img)
                self._spacing(img, after=80)

    def add_compensation(self, m):
        self.doc.add_page_break()

        title = m.get("title", "")
        last = m["client_last_name"]

        head = self.doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_font(head.add_run("CLAIM FOR COMPENSATION"), size=16, bold=True, underline=True)
        self._spacing(head, after=80)

        self._left_para(
            f"{title}. {last} is entitled to full and fair compensation for losses incurred "
            "in this accident. Relevant California law provides as follows:"
        )

        self._center_underline("Civil Code § 3333.  Torts in General")

        self._center_para(
            "Breach of obligation other than contract. For the breach of an obligation not "
            "arising from contract, the measure of damages, except where otherwise expressly "
            "provided by this code, is the amount which will compensate for all the detriment "
            "proximately caused thereby, whether it could have been anticipated or not."
        )

        self._left_para(
            f"{title}. {last} is also entitled to recover for all the pain and suffering "
            "experienced as a result of this accident:"
        )

        self._center_para(
            "Pain and suffering is a unitary concept, encompassing all of the physical and "
            "emotional trauma occasioned by an injury. Plaintiff is entitled to compensatory "
            "damages for all physical pain suffered . . . and also for resulting fright, "
            "nervousness, anxiety, worry, mortification, shock, humiliation, indignity, "
            "embarrassment, apprehension, terror, or ordeal. "
            "Capelouto v. Kaiser Foundation Hospitals (1972) 7 Cal.App.3d 889, 893-894."
        )

        self._left_para("The applicable jury instruction provides, in pertinent part, as follows:")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_font(p.add_run("CACI 3905A"), size=16, bold=True)
        self._spacing(p, after=80)

        self._center_para(
            f"{title}. {last} is also entitled to recover for all past and future noneconomic "
            "damages caused by your insured’s negligence, including physical pain, mental "
            "suffering, loss of enjoyment of life, disfigurement, physical impairment, "
            "inconvenience, grief, anxiety, humiliation, and emotional distress. "
            "(CACI 3905A, Physical Pain, Mental Suffering, and Emotional Distress)"
        )

    # -------------------------------------------------

    def _left_para(self, text):
        p = self.doc.add_paragraph()
        self._set_font(p.add_run(text))
        self._spacing(p, after=80)

    def _center_para(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_font(p.add_run(text))
        self._spacing(p, after=80)

    def _center_underline(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_font(p.add_run(text), underline=True)
        self._spacing(p, after=80)

    # -------------------------------------------------

    def generate(self, m, output):
        self.add_logo_and_date(m)
        self.add_insurance_block(m)
        self.add_title(m)
        self.add_client_info(m)
        self.add_settlement_notice()
        self.add_salutation_and_intro(m)
        self.add_incident_and_medical(m)
        self.add_compensation(m)

        self.doc.save(output)
        return output


# =====================================================
# MAIN
# =====================================================

if __name__ == "__main__":
    metadata = load_metadata(
        r"C:\Users\hp\Documents\pdf-extraction-pipeline\demandpdf_generation\Demand_metadata1.json"
    )

    generator = DemandNoteGenerator()
    output = generator.generate(metadata, "demand_note.docx")

    print(f"✅ Demand note created: {output}")
